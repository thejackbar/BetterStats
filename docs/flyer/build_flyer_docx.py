"""Build the BetterCricket flyer as an editable Word document.

Rebuilds the two-page PDF flyer as real Word content -- paragraphs, runs and
tables -- so the copy can be edited and the file exported back to PDF.

The dark page colour is painted TWICE on purpose: as the document background
(what Word shows on screen, and what LibreOffice exports) and as a full-page
VML rectangle anchored behind the text in the page header (what Word actually
prints, since Word leaves the document background out of print and PDF export
unless the reader has switched that option on).

    python3 docs/flyer/build_flyer_docx.py [out.docx]
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent

# ---------------------------------------------------------------- palette ---
BG      = "0B1220"   # page
CARD    = "101828"   # card fill
TINT    = "0A1D26"   # accent-tinted card fill
HAIR    = "1D2331"   # borders and rules
ACCENT  = "16C784"
BRIGHT  = "E6E8EF"   # headings
BODY    = "A3AAB9"   # body copy
DIM     = "8A90A2"   # captions
FAINT   = "5F6577"   # small print
INK     = "06231A"   # text on an accent fill

FONT = "Arial"

# Page geometry, in points, taken from the PDF.
PAGE_W, PAGE_H = 595.92, 842.88
MARGIN = 42.75
CONTENT = PAGE_W - 2 * MARGIN          # 510.4pt


# ------------------------------------------------------------- xml helpers --
# WordprocessingML property elements are a strict sequence, and both Word and
# LibreOffice refuse a file that puts them in the wrong order, so every insert
# goes through _ordered_insert with its element's own sequence.
TBL_PR_ORDER = (
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "tblJc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
)
TC_PR_ORDER = (
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
)
P_PR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
)
R_PR_ORDER = (
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
    "dstrike", "outline", "shadow", "emboss", "imprint", "noProof",
    "snapToGrid", "vanish", "webHidden", "color", "spacing", "w", "kern",
    "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd",
    "fitText", "vertAlign", "rtl", "cs", "em", "lang",
)


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def _local(tag):
    return tag.split("}")[-1]


def _ordered_insert(parent, element, order):
    """Insert element into parent at its schema position, replacing any twin."""
    name = _local(element.tag)
    for existing in parent.findall(qn("w:" + name)):
        parent.remove(existing)
    rank = order.index(name)
    for child in parent:
        child_name = _local(child.tag)
        if child_name not in order or order.index(child_name) > rank:
            child.addprevious(element)
            return element
    parent.append(element)
    return element


def shade(props, fill, order):
    _ordered_insert(props, _el("w:shd", val="clear", color="auto", fill=fill), order)


def cell_bg(cell, fill):
    shade(cell._tc.get_or_add_tcPr(), fill, TC_PR_ORDER)


def cell_borders(cell, color=HAIR, size=6, sides=("top", "left", "bottom", "right")):
    """size is in eighths of a point (6 = 0.75pt, matching the PDF)."""
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        borders.append(
            _el("w:" + side, val="single" if side in sides else "nil",
                sz=size if side in sides else 0, space=0, color=color)
        )
    _ordered_insert(cell._tc.get_or_add_tcPr(), borders, TC_PR_ORDER)


def cell_pad(cell, top=6, left=8, bottom=6, right=8):
    """Padding in points."""
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el("w:" + side, w=int(val * 20), type="dxa"))
    _ordered_insert(cell._tc.get_or_add_tcPr(), mar, TC_PR_ORDER)


def cell_width(cell, pts):
    _ordered_insert(cell._tc.get_or_add_tcPr(),
                    _el("w:tcW", w=int(pts * 20), type="dxa"), TC_PR_ORDER)


def drop(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def clear(container):
    """Empty a cell of the placeholder paragraph the grid seeded it with."""
    for paragraph in list(container.paragraphs):
        drop(paragraph)
    return container


# ------------------------------------------------------------ text helpers --
def run(paragraph, text, size=7.4, color=BODY, bold=False, italic=False,
        track=None, highlight=None):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    rPr = r._r.get_or_add_rPr()
    # East-Asian / complex-script font names, so Word does not substitute.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _ordered_insert(rPr, _el("w:rFonts"), R_PR_ORDER)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn("w:" + attr), FONT)
    if track:                                   # letter spacing, in points
        _ordered_insert(rPr, _el("w:spacing", val=int(track * 20)), R_PR_ORDER)
    if highlight:                               # run-level background fill
        shade(rPr, highlight, R_PR_ORDER)
    return r


def para(container, space_before=0, space_after=0, line=None, align=None,
         keep_with_next=False, page_break=False):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)
    if align is not None:
        pf.alignment = align
    pf.keep_with_next = keep_with_next
    pf.page_break_before = page_break
    return p


def text_para(container, text, size=7.4, color=BODY, bold=False, italic=False,
              line=11.2, space_before=0, space_after=0, align=None, track=None,
              keep_with_next=False):
    p = para(container, space_before, space_after, line, align, keep_with_next)
    run(p, text, size, color, bold, italic, track)
    return p


def rule(container, color=HAIR, space_before=6, space_after=6, size=6):
    """A hairline rule drawn as a paragraph bottom border."""
    p = para(container, space_before, space_after, line=1)
    borders = OxmlElement("w:pBdr")
    borders.append(_el("w:bottom", val="single", sz=size, space=1, color=color))
    _ordered_insert(p._p.get_or_add_pPr(), borders, P_PR_ORDER)
    run(p, "", 1)
    return p


# ---------------------------------------------------------------- tables ----
def grid(doc, widths, gap, rows=1):
    """A fixed-layout, borderless table: content columns separated by spacers.

    Returns the list of content cells.
    """
    n = len(widths)
    cols = []
    for i, w in enumerate(widths):
        cols.append(w)
        if i < n - 1:
            cols.append(gap)

    table = doc.add_table(rows=rows, cols=len(cols))
    table.autofit = False

    # LibreOffice and Word both size a fixed-layout table from w:tblGrid, not
    # from the per-cell widths, so the grid has to be rewritten or the columns
    # come out evenly divided.
    grid_el = table._tbl.find(qn("w:tblGrid"))
    if grid_el is not None:
        table._tbl.remove(grid_el)
    grid_el = OxmlElement("w:tblGrid")
    for w in cols:
        grid_el.append(_el("w:gridCol", w=int(w * 20)))
    table._tbl.insert(list(table._tbl).index(table._tbl.tblPr) + 1, grid_el)

    tblPr = table._tbl.tblPr
    _ordered_insert(tblPr, _el("w:tblW", w=int(sum(cols) * 20), type="dxa"), TBL_PR_ORDER)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el("w:" + side, val="nil"))
    _ordered_insert(tblPr, borders, TBL_PR_ORDER)
    _ordered_insert(tblPr, _el("w:tblInd", w=0, type="dxa"), TBL_PR_ORDER)
    _ordered_insert(tblPr, _el("w:tblLayout", type="fixed"), TBL_PR_ORDER)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        mar.append(_el("w:" + side, w=0, type="dxa"))
    _ordered_insert(tblPr, mar, TBL_PR_ORDER)

    out = []
    for row in table.rows:
        for cell, w in zip(row.cells, cols):
            cell_width(cell, w)
            cell_pad(cell, 0, 0, 0, 0)
            pf = cell.paragraphs[0].paragraph_format
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(1)
            run(cell.paragraphs[0], "", 1)
        out.append([row.cells[i * 2] for i in range(n)])
    return out[0] if rows == 1 else out


def card(cell, fill=CARD, border=HAIR, pad=(7, 9, 7, 9), border_sides=("top", "left", "bottom", "right"), border_size=6):
    cell_bg(cell, fill)
    cell_borders(cell, border, border_size, border_sides)
    cell_pad(cell, *pad)
    clear(cell)   # drop the placeholder paragraph the grid seeded
    return cell


def spacer(doc, height=6):
    para(doc, 0, 0, line=height)


# ------------------------------------------------------- page background ----
def paint_background(doc):
    """Dark page colour: document background + a printable header rectangle."""
    body = doc.element.body
    background = OxmlElement("w:background")
    background.set(qn("w:color"), BG)
    doc.element.insert(0, background)

    settings = doc.settings.element
    display = OxmlElement("w:displayBackgroundShape")
    settings.append(display)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(1)
        r = p.add_run()
        r.font.size = Pt(1)
        style = (
            "position:absolute;margin-left:0;margin-top:0;"
            f"width:{PAGE_W}pt;height:{PAGE_H}pt;z-index:-251658752;"
            "mso-position-horizontal:left;mso-position-horizontal-relative:page;"
            "mso-position-vertical:top;mso-position-vertical-relative:page"
        )
        pict = parse_xml(
            f'<w:pict {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml">'
            f'<v:rect style="{style}" fillcolor="#{BG}" stroked="f"/>'
            "</w:pict>"
        )
        r._r.append(pict)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        def field(instr):
            r = p.add_run()
            r.font.name = FONT
            r.font.size = Pt(6)
            r.font.color.rgb = RGBColor.from_string("4D5364")
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instr_el = OxmlElement("w:instrText")
            instr_el.set(qn("xml:space"), "preserve")
            instr_el.text = instr
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            r._r.append(begin)
            r._r.append(instr_el)
            r._r.append(end)

        field(" PAGE ")
        run(p, " / ", 6, "4D5364")
        field(" NUMPAGES ")


# ------------------------------------------------------------- components ---
def masthead(doc, tagline, contact, page_break=False):
    if page_break:
        # The break sits on its own paragraph rather than inside the table:
        # page-break-before on a paragraph in a cell leaves a blank page behind.
        para(doc, 0, 0, line=1, page_break=True)

    cells = grid(doc, [250, CONTENT - 250 - 8], 8)
    left, right = cells

    clear(left)
    # The exact line has to clear the logo, or its top is clipped.
    p = para(left, 0, 0, line=24)
    logo = HERE / "logo.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), height=Pt(22))
    run(p, "  ", 15.8, BRIGHT)
    run(p, "Better", 15.8, BRIGHT, bold=True)
    run(p, "Cricket", 15.8, ACCENT, bold=True)

    clear(right)
    text_para(right, tagline, 7.6, DIM, line=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    text_para(right, contact, 7.9, ACCENT, bold=True, line=11,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

    rule(doc, HAIR, space_before=8, space_after=0)


def eyebrow(doc, text, space_before=14):
    text_para(doc, text, 7.4, ACCENT, bold=True, line=10,
              space_before=space_before, space_after=2, track=1.1,
              keep_with_next=True)


def bullet(container, text, size=7.3, line=10.2, space_before=1.5):
    """A bullet whose wrapped lines run flush left, not indented under the text."""
    p = para(container, space_before, 0, line)
    run(p, "·  ", size, ACCENT, bold=True)
    run(p, text, size, BODY)
    return p


def module_card(cell, name, price, promise, bullets, width, pill=None,
                tinted=False, columns=1):
    """One module: what it is, then what the club actually gets."""
    card(cell, fill=TINT if tinted else CARD, border=ACCENT if tinted else HAIR,
         pad=(7, 9, 7, 9))
    inner = width - 18
    name_w = min(200, inner - 62)
    head = grid(cell, [name_w, inner - name_w - 6], 6)
    clear(head[0])
    p = para(head[0], 0, 0, line=11)
    run(p, name, 8.5, BRIGHT, bold=True)
    if pill:
        run(p, "  ", 8.5, BRIGHT)
        run(p, "  " + pill + "  ", 6.4, INK, bold=True, track=0.5, highlight=ACCENT)
    clear(head[1])
    text_para(head[1], price, 8.1, ACCENT, bold=True, line=11,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
    if promise:
        text_para(cell, promise, 7.6, ACCENT, bold=True, line=10.5, space_before=3)
    if columns == 1:
        for b in bullets:
            bullet(cell, b)
        return
    # Two columns: fill the left one first, so a short list stays left-heavy
    # and the card never ends on a lone orphan in the right column.
    half = (len(bullets) + 1) // 2
    col_w = (inner - 12) / 2
    cols = grid(cell, [col_w, col_w], 12)
    for column, items in zip(cols, (bullets[:half], bullets[half:])):
        clear(column)
        for i, b in enumerate(items):
            bullet(column, b, space_before=0 if i == 0 else 1.5)


# ------------------------------------------------------------------ page 1 --
def page_one(doc):
    masthead(doc, "The platform Australian cricket clubs run on", "betterat.cricket")

    p = para(doc, space_before=13, space_after=8, line=30)
    run(p, "We do cricket", 26.2, BRIGHT, bold=True)
    run(p, "… ", 26.2, DIM, bold=True)
    run(p, "Better.", 26.2, ACCENT, bold=True)

    p = para(doc, space_after=11, line=13.6)
    run(p, "Every club runs the same way: one volunteer's spreadsheet, a website "
           "nobody can edit, a design app, a mail tool and a group chat that ends "
           "up picking the side. Nothing talks to anything else, so the season gets "
           "typed in five times. ", 8.6, BODY)
    run(p, "BetterCricket is one platform for the whole club, fed by one match feed",
        8.6, BRIGHT, bold=True)
    run(p, ". Enter the season once and the stats, the site, the posts, the money "
           "and the paperwork all run off it.", 8.6, BODY)

    stats = [
        ("Under 1 hr", "From first sync to a live, branded club site"),
        ("Overnight", "Each round's results are in by Sunday morning"),
        ("Decades", "Of history brought in, with no migration fee"),
        ("14 days", "Free trial of every module, no card needed"),
    ]
    cells = grid(doc, [122.3] * 4, 6.7)
    for cell, (head, note) in zip(cells, stats):
        card(cell, pad=(6, 8, 6, 8))
        text_para(cell, head, 11.6, ACCENT, bold=True, line=13, space_after=3)
        text_para(cell, note, 6.9, DIM, line=9)

    spacer(doc, 13)

    eyebrow(doc, "THE JOBS NOBODY VOLUNTEERED FOR", space_before=0)
    text_para(doc, "Every one of these came from a club secretary telling us how "
                   "their weekend actually goes.", 8.0, DIM, line=11, space_after=5,
              keep_with_next=True)

    jobs = [
        ("01", "A history nobody can lay their hands on",
         "One career split across three spellings, fifty years in a filing cabinet, "
         "and a premiership side nobody can name. The records live in whatever the "
         "last secretary used, and they leave when he does."),
        ("02", "Saturday's \"who's in?\"",
         "Availability arrives by group chat, gets copied into a spreadsheet, and "
         "the side gets picked twice because two grades both wanted the same "
         "all-rounder."),
        ("03", "Five tools and five bills",
         "A spreadsheet, a site builder, Canva, a mail tool and a payments app, none "
         "of which talk to each other. Nobody can change the website except the one "
         "bloke who built it."),
        ("04", "Money tracked in three places",
         "Subs here, match fees there, canteen takings in a tin. Chased by text, "
         "reconciled by hand, and the treasurer finds out in June who never paid."),
        ("05", "Compliance you hear about late",
         "A working-with-children check that lapsed in March, an RSA nobody holds, "
         "and a ground inspection that was due before the first game."),
        ("06", "The committee's memory",
         "Motions carried and never written up, actions nobody owns, and a strategic "
         "plan that stalls the day its author steps down."),
    ]
    job_rows = grid(doc, [251, 251], 8, rows=3)
    for pair, cells in zip((jobs[0:2], jobs[2:4], jobs[4:6]), job_rows):
        for cell, (num, title, body) in zip(cells, pair):
            cell_pad(cell, 0, 0, 6, 0)
            clear(cell)
            inner = grid(cell, [16, 235], 0)
            clear(inner[0])
            text_para(inner[0], num, 7.4, ACCENT, bold=True, line=11)
            clear(inner[1])
            text_para(inner[1], title, 7.9, BRIGHT, bold=True, line=11)
            text_para(inner[1], body, 7.3, BODY, line=10.4, space_before=2)

    eyebrow(doc, "INCLUDED WITH EVERY CLUB", space_before=10)
    text_para(doc, "BetterStats is the base every club starts on, and on its own it "
                   "replaces the spreadsheet, the filing cabinet and the club site.",
              8.0, DIM, line=11, space_after=5, keep_with_next=True)

    module_card(
        grid(doc, [CONTENT], 0)[0], "BetterStats", "$399 / year",
        "Your club's whole history, public, current and finally in one place.",
        ["Every batting, bowling and fielding figure the club holds, reconciled "
         "across decades and kept current on its own",
         "Scorecards, fixtures and ladders going back as far as your records do",
         "Duplicate players found and merged in a click, keeping every innings, "
         "spell and catch",
         "Photograph an old scorebook page and the figures lift off it, checked "
         "against what you already hold",
         "No migration fee and no cap on seasons, however far back the club goes",
         "Your own public site, at your own address, in your colours and crest",
         "A profile for every player, leaderboards, all-time and partnership "
         "records, and the honour board",
         "Head-to-head splits: how a player goes against every club, ground and "
         "format they have played",
         "Filter any of it by grade type and match type, so Under-14s never pad a "
         "senior average",
         "StatLab, so any question the committee asks gets a real answer and a "
         "report you can save",
         "Milestones flagged before they happen, and awards and honours recorded "
         "against the player",
         "A season yearbook that writes itself, and Club Room Mode for the TV in "
         "the clubhouse"],
        CONTENT, pill="INCLUDED", tinted=True, columns=2)

    spacer(doc, 14)

    eyebrow(doc, "ONE PLATFORM, NOT FIVE SUBSCRIPTIONS", space_before=0)
    band = card(grid(doc, [CONTENT], 0)[0], fill=TINT, border=ACCENT,
                pad=(8, 10, 8, 10), border_sides=("left",), border_size=11)
    p = para(band, 0, 0, line=12)
    run(p, "Enter the season once.  ", 8.6, ACCENT, bold=True)
    run(p, "The scorecard that moves a player's average is the same one that fills "
           "a match-day post, prices a fantasy round, shows the selectors who is in "
           "form, books the match fee against the right member and turns up in the "
           "yearbook. Nothing is typed twice, nothing drifts out of step, and the "
           "club gets one bill instead of five.", 7.8, BODY)

    spacer(doc, 8)

    quote = card(grid(doc, [CONTENT], 0)[0], fill=BG, border=ACCENT,
                 pad=(5, 12, 5, 8), border_sides=("left",), border_size=11)
    text_para(quote, "\"At last we have a complete stats package that lets us view "
                     "the club's entire history across every statistic imaginable. "
                     "It's made pretty much every spreadsheet we had redundant, and "
                     "we had a lot.\"", 8.1, BRIGHT, italic=True, line=12)
    text_para(quote, "Tristram Fletcher · Secretary, Applecross Cricket Club",
              7.2, DIM, line=11, space_before=4)


# ------------------------------------------------------------------ page 2 --
def page_two(doc):
    masthead(doc, "The modules, what they cost, and how to start",
             "support@bettersports.com.au", page_break=True)

    eyebrow(doc, "ADD ONLY WHAT YOU NEED", space_before=11)
    text_para(doc, "Each module is a job the club is already doing by hand. Take "
                   "one, take the lot, or start with none and add them mid-season.",
              8.0, DIM, line=11, space_after=5, keep_with_next=True)

    pair_w = (CONTENT - 6) / 2
    row = grid(doc, [pair_w, pair_w], 6)
    module_card(
        row[0], "BetterSelect", "$149 / yr",
        "Pick the side without the Thursday group chat.",
        ["Players set their own availability from a link. No app, no account, "
         "nothing to install",
         "Every squad and grade on one board, so two sides can't pick the same "
         "all-rounder",
         "Name the XI on a numbered batting order with each player's recent form "
         "beside them",
         "Association rules checked as you pick: age limits, overseas caps, junior "
         "bowling workloads, finals qualification",
         "Nets off a QR check-in and a rotation timer that keeps the queue moving",
         "3-2-1 votes through the same link, counted for you and revealed on screen "
         "at the awards night"],
        pair_w)
    module_card(
        row[1], "BetterSocials", "$149 / yr",
        "Your website and match-day posts, both fed by your scorecards.",
        ["A full club website: news, pages, galleries, sponsors and honour boards, "
         "on top of the stats pages every club gets",
         "No hosting bill, no site-builder subscription, and no waiting on the one "
         "volunteer who knows how it works",
         "A post designer built for cricket: blocks on a canvas, layers, undo and "
         "redo, carousels and a club photo library",
         "Drop in a match link and the card fills itself with the real figures, in "
         "your colours",
         "Live blocks for fixtures, results or a player's career, so a post is never "
         "out of date"],
        pair_w)

    spacer(doc, 6)

    wide = grid(doc, [CONTENT], 0)[0]
    module_card(
        wide, "BetterAdmin", "$149 / yr",
        "The back office. Say goodbye to spreadsheets, Mailchimp and more.",
        ["Subs, memberships and match fees that settle themselves as payments land, "
         "with who is financial and who owes on one screen",
         "Square for card and canteen takings, Xero for the books, so the treasurer "
         "stops re-keying the same figures",
         "Stock, canteen and merch, with low-stock alerts and what each member still "
         "owes on their kit",
         "A sponsor and grant pipeline, so the money conversations don't sit in one "
         "person's inbox",
         "One directory of everyone at the club: players, parents, volunteers, "
         "officials and life members",
         "Bulk email to that same list, with segments and templates, in place of a "
         "separate mail subscription",
         "The volunteer roster, with hours logged and ready for a grant application",
         "Working-with-children, RSA and first-aid tickets tracked, with warnings "
         "before they lapse",
         "Committee meetings with agendas, minutes, motions and actions, tied to the "
         "club's strategic plan",
         "A club diary for the annual jobs: registrations, insurance, ground "
         "bookings, the AGM",
         "Events and ticketing for the presentation night and the fundraiser",
         "Facilities and assets: bookings, hire, maintenance history and replacement "
         "dates"],
        CONTENT, columns=2)

    spacer(doc, 6)

    row = grid(doc, [pair_w, pair_w], 6)
    module_card(
        row[0], "BetterIQ", "$249 / yr",
        "Know exactly how to beat your opposition before the toss.",
        ["A dossier on any upcoming opponent, built from your own scorecards without "
         "anyone requesting it",
         "Their danger players named, with a printable captain's cheat sheet",
         "A best-available XI, player trends and milestone forecasts",
         "Your own side analysed too: partnerships, collapses, par scores and who "
         "actually wins you games"],
        pair_w)
    module_card(
        row[1], "BetterFantasyCricket", "$49 / yr",
        "The comp that gets the whole club talking and engaged.",
        ["Salary cap or draft, scored off your club's real scorecards across every "
         "grade",
         "Squads of 12, the captain scores double, and the best 11 count each round",
         "A club ladder plus private mini-leagues, and members join from a link",
         "Free to enter, so parents and supporters end up following the 4ths as "
         "closely as the 1sts"],
        pair_w)

    eyebrow(doc, "WHAT IT COSTS", space_before=11)

    money = card(grid(doc, [CONTENT], 0)[0], fill=CARD, border=HAIR, pad=(8, 10, 8, 10))
    p = para(money, 0, 0, line=12)
    for i, (label, price) in enumerate([
        ("BetterStats", "$399"), ("BetterSelect", "$149"), ("BetterSocials", "$149"),
        ("BetterAdmin", "$149"), ("BetterIQ", "$249"), ("BetterFantasyCricket", "$49"),
    ]):
        if i:
            run(p, "   ·   ", 8.0, FAINT)
        run(p, label + " ", 8.0, BODY)
        run(p, price, 8.0, BRIGHT, bold=True)
    p = para(money, 5, 0, line=12)
    run(p, "Bundle the four modules with BetterStats and it is ", 8.4, BODY)
    run(p, "$949 a year", 8.4, ACCENT, bold=True)
    run(p, ". Everything, BetterFantasyCricket included, is ", 8.4, BODY)
    run(p, "$998", 8.4, ACCENT, bold=True)
    run(p, ".", 8.4, BODY)
    text_para(money, "An annual licence, billed once and paid by card through "
                     "Stripe. Flat per club: one team or fifty, juniors and seniors, "
                     "men's and women's, with unlimited players and seasons. About "
                     "$955 less than the five tools it replaces, and the historical "
                     "import a rival charges $499 to $1,000 for is included.",
              7.2, DIM, line=10.5, space_before=5)

    rule(doc, HAIR, space_before=9, space_after=8)

    cells = grid(doc, [380, CONTENT - 380 - 10], 10)
    left = cells[0]
    clear(left)
    text_para(left, "Get your club on BetterCricket.", 12.4, BRIGHT, bold=True,
              line=15, space_after=4)
    p = para(left, 0, 0, line=11.5)
    run(p, "Start the 14-day free trial of every module at ", 7.8, DIM)
    run(p, "betterat.cricket/trial", 7.8, ACCENT, bold=True)
    run(p, ", or email us and we'll walk your committee through it on a "
           "15-minute call. Setting up takes about half an hour: we run your first "
           "sync, we tidy the history with you, and your club goes public.",
        7.8, DIM)

    button = cells[1]
    clear(button)
    inner = grid(button, [110], 0)[0]
    card(inner, fill=ACCENT, border=ACCENT, pad=(7, 8, 7, 8))
    text_para(inner, "Start the free trial →", 8.1, INK, bold=True, line=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    text_para(doc, "BetterCricket is the cricket platform from BetterSports · "
                   "ABN 32 624 335 397 · betterat.cricket · "
                   "support@bettersports.com.au", 6.6, FAINT, line=10,
              space_before=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def close_cells(doc):
    """A table cell must end with a paragraph.

    Several cells here hold nothing but a nested table (the numbered jobs, the
    call-to-action button). OOXML requires a trailing w:p, and a file without
    one is rejected outright rather than merely rendered oddly, so give each
    such cell a 1pt paragraph that takes up no visible room.
    """
    for tc in doc.element.body.iter(qn("w:tc")):
        children = list(tc)
        if children and _local(children[-1].tag) != "p":
            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            spacing = _el("w:spacing", after=0, line=20, lineRule="exact")
            pPr.append(spacing)
            rPr = OxmlElement("w:rPr")
            rPr.append(_el("w:sz", val=2))
            pPr.append(rPr)
            p.append(pPr)
            tc.append(p)


# -------------------------------------------------------------------- main --
def build(out_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Pt(PAGE_W)
    section.page_height = Pt(PAGE_H)
    section.left_margin = Pt(MARGIN)
    section.right_margin = Pt(MARGIN)
    section.top_margin = Pt(34)
    section.bottom_margin = Pt(25)
    section.header_distance = Pt(14)
    section.footer_distance = Pt(14)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(7.4)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_after = Pt(0)

    paint_background(doc)
    add_footer(doc)

    # The default template may open with an empty paragraph; the flyer supplies
    # its own content, so drop it.
    if doc.paragraphs:
        first = doc.paragraphs[0]
        first._p.getparent().remove(first._p)

    page_one(doc)
    page_two(doc)
    close_cells(doc)

    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else str(HERE / "BetterCricket_Flyer.docx")
    build(out)
